from docx import Document
import xlrd
from xlutils.copy import copy
import os
row = 0
row2 = 0


def table(path_file):
    doc = Document(path_file)
    tables = doc.tables
    alist1 = []
    alist2 = []
    alist3 = []
    alist4 = []
    for i in range(20):  # 表格个数
        try:
            table = tables[i]
            for row_n, row in enumerate(table.rows):
                column = len(row.cells)  # 列宽
                data_docx = doc.tables[i].cell(row_n, 0).text
                # data_docx = data_docx.replace('\n', '')
                # data_docx = data_docx.strip()
                a = data_docx.split()
                data_docx = ''.join(a)
                # print(data_docx)
                if (data_docx == '时间') or (data_docx == '日期'):
                    for c in range(1, column):
                        sj = doc.tables[i].cell(row_n, c).text
                        sj = sj.replace('\n', '')
                        alist1.append(sj)
                if '诊疗' in data_docx :
                    for c in range(1, column):
                        zl = doc.tables[i].cell(row_n, c).text
                        alist2.append(zl)
                if '医嘱'in data_docx :
                    for c in range(1, column):
                        yz = doc.tables[i].cell(row_n, c).text
                        alist3.append(yz)
                    # print(alist3)
                if '护理' in data_docx:
                    for c in range(1, column):
                        hl = doc.tables[i].cell(row_n, c).text
                        alist4.append(hl)
        except IndexError:
            break
    return alist1, alist2, alist3, alist4


def search_file(target):
    os.chdir(target)
    files_name = os.walk(target)
    for file_name in files_name:
        for each in file_name[2]:
            each_path = file_name[0] + os.sep + each
            if ('~' or '$') in each:
                continue
            if os.path.splitext(each)[1] == '.docx':
                data = table(each_path)
                write_sj('D:\\pythoncode\\4.表格与文档操作\\操作word的表格\\数据1.xls', each_path, data[0])
                write('D:\\pythoncode\\4.表格与文档操作\\操作word的表格\\数据2.xls', each_path, data[1], 'ZL')
                write('D:\\pythoncode\\4.表格与文档操作\\操作word的表格\\数据2.xls', each_path, data[2], 'YZ')
                write('D:\\pythoncode\\4.表格与文档操作\\操作word的表格\\数据2.xls', each_path, data[3], 'HL')
                del data
            else:
                with open('D:\\pythoncode\\4.表格与文档操作\\操作word的表格\\未写入.txt', 'a') as f:
                    f.writelines(each_path + '\n')
# 1806000001	R0001	YZ	2	1	儿科护理常规	1	0	NULL	NULL	NULL
# 1806000001	R0002	住院第2-13天	0	2	0	NULL	NULL	NULL	1


def write_sj(file, source, data):
    global row2
    f = file
    rb = xlrd.open_workbook(f, formatting_info=True)
    wb = copy(rb)
    sheet = wb.get_sheet(0)
    n = 0
    for each in data:
        n += 1
        each = each.replace('\n', '')
        each = each.strip()
        if each == '':
            continue
        sheet.write(row2, 0, source)
        R = 'R00' + str(n)
        sheet.write(row2, 1, R)
        sheet.write(row2, 2, each)
        sheet.write(row2, 3, 0)
        sheet.write(row2, 4, n)
        sheet.write(row2, 5, 0)
        row2 += 1
    wb.save(f)


def write(file, source, data, lx):
    global row
    f = file
    rb = xlrd.open_workbook(f, formatting_info=True)
    wb = copy(rb)
    sheet = wb.get_sheet(0)
    n = 0
    for each in data:
        s = each.split('\n')
        xh = 1
        n += 1
        for each_s in s:
            each_s = each_s.replace('□', '')
            each_s = each_s.strip()
            if each_s == '':
                continue
            sheet.write(row, 5, each_s)
            sheet.write(row, 3, xh)
            sheet.write(row, 0, source)
            R = 'R00' + str(n)
            sheet.write(row, 1, R)
            sheet.write(row, 2, lx)
            sheet.write(row, 4, 1)
            sheet.write(row, 6, 1)
            sheet.write(row, 7, 0)
            xh += 1
            row += 1
    wb.save(f)


target = r"E:\DOCX文件"
search_file(target)
