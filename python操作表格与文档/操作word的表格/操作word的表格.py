from docx import Document
import os


#  读取一个word的所有数据
def table(path_file):
    file_name = os.path.basename(path_file)
    xh = 0  # 行号
    doc = Document(path_file)
    tables = doc.tables
    data_list = []
    for i in range(10):  # 表格个数
        try:
            table = tables[i]
            for row_n, row in enumerate(table.rows):
                xh += 1
                column = len(row.cells)  # 列宽
                for c in range(column):
                    data_docx = doc.tables[i].cell(row_n, c).text
                    if data_docx == '':  # 单元格去空格
                        continue
                    data = (data_docx, row_n + 1, c + 1, i + 1, xh, file_name)
                    data_list.append(data)
        except IndexError:
            break
    return data_list


path_file = "E:\\2017年拟发布临床路径病种表单202种\\血液科（11个）\\县级\\初治原发免疫性血小板减少症（县医院适用版）.docx"
d = table(path_file)
for each in d:
    print(each)



